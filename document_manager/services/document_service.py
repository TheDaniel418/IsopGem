"""
Purpose: Provides document management business logic for the application.

This file is part of the document_manager pillar and serves as a service component.
It is responsible for handling document operations, text extraction, and processing.

Key components:
- DocumentService: Service class for document management operations

Dependencies:
- document_manager.models.document: For Document model
- document_manager.repositories.document_repository: For document persistence
- PyMuPDF: For PDF processing (imported as fitz)
- python-docx: For DOCX processing
"""

import concurrent.futures
import os
import re
import shutil
from datetime import datetime
from glob import glob
from pathlib import Path
from typing import Any, Dict, List, Optional, Tuple, Union

import fitz  # PyMuPDF
from docx import Document as DocxDocument
from loguru import logger

from document_manager.models.document import Document, DocumentType
from document_manager.repositories.document_repository import DocumentRepository


class DocumentService:
    """Service for document management operations."""

    # Document types for OpenDocument formats
    ODT = "odt"  # OpenDocument Text
    ODS = "ods"  # OpenDocument Spreadsheet
    ODP = "odp"  # OpenDocument Presentation

    # Windows Symbol font to Unicode Greek letter mapping
    SYMBOL_TO_GREEK = {
        "a": "α",
        "b": "β",
        "g": "γ",
        "d": "δ",
        "e": "ε",
        "z": "ζ",
        "h": "η",
        "q": "θ",
        "i": "ι",
        "k": "κ",
        "l": "λ",
        "m": "μ",
        "n": "ν",
        "x": "ξ",
        "o": "ο",
        "p": "π",
        "r": "ρ",
        "s": "σ",
        "ς": "ς",
        "t": "τ",
        "u": "υ",
        "f": "φ",
        "c": "χ",
        "y": "ψ",
        "w": "ω",
        "A": "Α",
        "B": "Β",
        "G": "Γ",
        "D": "Δ",
        "E": "Ε",
        "Z": "Ζ",
        "H": "Η",
        "Q": "Θ",
        "I": "Ι",
        "K": "Κ",
        "L": "Λ",
        "M": "Μ",
        "N": "Ν",
        "X": "Ξ",
        "O": "Ο",
        "P": "Π",
        "R": "Ρ",
        "S": "Σ",
        "T": "Τ",
        "U": "Υ",
        "F": "Φ",
        "C": "Χ",
        "Y": "Ψ",
        "W": "Ω",
    }

    # Special Symbol font characters that map to other Greek characters or symbols
    SYMBOL_SPECIAL = {
        # Symbol font uses "V" for the end-of-word sigma (ς)
        "V": "ς",
        # Other special mappings for mathematical symbols often used in Greek texts
        "@": "∈",  # Element of
        "∃": "∃",  # There exists
        "∀": "∀",  # For all
        "∂": "∂",  # Partial differential
        "∇": "∇",  # Nabla/Del
        "∞": "∞",  # Infinity
        "∝": "∝",  # Proportional to
    }

    def __init__(self, document_repository: Optional[DocumentRepository] = None):
        """Initialize the document service.

        Args:
            document_repository: Repository for document storage, created if not provided
        """
        self.document_repository = document_repository or DocumentRepository()
        self.storage_dir = Path("data/documents")

        # Ensure storage directory exists
        os.makedirs(self.storage_dir, exist_ok=True)

    def _convert_symbol_to_greek(self, text: str) -> Tuple[str, bool]:
        """Convert Symbol font characters to proper Unicode Greek letters.

        This function detects text that appears to be in Symbol font (common in older documents)
        and converts it to proper Unicode Greek characters for better searchability.

        Args:
            text: Text that may contain Symbol font encoded Greek letters

        Returns:
            Tuple of (converted_text, was_converted)
        """
        # MUCH MORE strict markers of Symbol font use
        # We need multiple strong indicators and explicit evidence before converting

        # Check first for explicit Symbol font declarations
        explicit_symbol_markers = [
            "Symbol font",
            "Symbol typeface",
            "Greek using Symbol",
            "Symbol encoding",
        ]

        has_explicit_marker = False
        for marker in explicit_symbol_markers:
            if marker.lower() in text.lower()[:1000]:
                logger.info(f"Found explicit Symbol font declaration: '{marker}'")
                has_explicit_marker = True
                break

        # Without explicit markers, look for very specific Greek letter sequences that would be nonsensical in English
        if not has_explicit_marker:
            # These are sequences that would make sense in Greek but are gibberish in English
            very_specific_greek_markers = [
                r"\balfa\s+kai\s+wmega\b",  # alpha and omega in Symbol encoding
                r"\bqeorem\s+[0-9]+\b",  # theorem with theta
                r"\bgamma\s+delta\b",
                r"\bfsi\s+fonction\b",  # psi function
                r"\bprwton\b",  # proton with omega
                r"\bqeta\b",  # theta spelled with Symbol
                r"\bs\^2\s*\+\s*p\^2\b",  # Mathematical formula with Symbol chars
            ]

            has_specific_sequences = False
            for marker in very_specific_greek_markers:
                if re.search(marker, text, re.IGNORECASE):
                    logger.info(
                        f"Found definite Greek sequence in Symbol font: '{marker}'"
                    )
                    has_specific_sequences = True
                    break

            if not has_specific_sequences:
                # No explicit markers or specific Greek sequences found
                # Only consider statistical measures as a last resort, with much higher thresholds

                # Sample text for analysis
                sample_text = text[:3000]

                # Define characters that are common in Symbol font for Greek encoding
                greek_specific_chars = set("abgdezhqiklmnxoprVsctufcyw")

                # Count sequences of 4+ consecutive Symbol-mappable characters
                # (Raising from 3 to 4 characters to be more strict)
                sequence_count = 0
                current_sequence = 0

                for char in sample_text:
                    if char.lower() in greek_specific_chars:
                        current_sequence += 1
                    else:
                        if (
                            current_sequence >= 4
                        ):  # More strict: require 4+ char sequences
                            sequence_count += 1
                        current_sequence = 0

                # Much more strict statistical criteria
                # 1. Need at least 10 Greek-like sequences (up from 5)
                # 2. At least 25% of text must be Symbol-mappable (up from 15%)
                # 3. English vowel ratio must be unusual (indicating non-English text)

                symbol_chars_count = sum(
                    1 for c in sample_text if c.lower() in greek_specific_chars
                )
                percentage_symbol = (
                    symbol_chars_count / len(sample_text) if sample_text else 0
                )

                # Check English vowel ratio (English is typically 35-45% vowels)
                english_vowels = set("aeiouy")
                vowel_count = sum(1 for c in sample_text.lower() if c in english_vowels)
                vowel_percentage = vowel_count / len(sample_text) if sample_text else 0

                # Vowel ratio is either very low or very high (non-English-like)
                unusual_vowel_ratio = vowel_percentage < 0.2 or vowel_percentage > 0.5

                if (
                    sequence_count >= 10
                    and percentage_symbol > 0.25
                    and unusual_vowel_ratio
                ):
                    logger.info(
                        f"Statistical analysis suggests Symbol font: {sequence_count} sequences, "
                        f"{percentage_symbol:.2%} Symbol chars, {vowel_percentage:.2%} vowels"
                    )
                    # Additional verification: check for common English words that would become nonsense in Greek
                    common_english_words = [
                        "the",
                        "and",
                        "for",
                        "that",
                        "this",
                        "with",
                        "from",
                    ]
                    english_word_count = 0
                    for word in common_english_words:
                        if re.search(r"\b" + word + r"\b", sample_text, re.IGNORECASE):
                            english_word_count += 1

                    # If we find multiple common English words, it's probably not Greek
                    if english_word_count >= 3:
                        logger.info(
                            f"Found {english_word_count} common English words, likely not Symbol font Greek"
                        )
                        return text, False
                else:
                    # Does not meet the stricter statistical criteria
                    return text, False

        # If we get here, we've decided this is Symbol font Greek
        # Make a copy of the text to modify
        converted_text = text

        # Replace common Symbol font encoded Greek with proper Unicode
        for latin, greek in self.SYMBOL_TO_GREEK.items():
            converted_text = converted_text.replace(latin, greek)

        # Handle special Symbol font characters
        for symbol, replacement in self.SYMBOL_SPECIAL.items():
            converted_text = converted_text.replace(symbol, replacement)

        # Analyze the impact of the conversion
        self._analyze_conversion_impact(text, converted_text)

        # Check for maximum conversion threshold (50%)
        # If more than 50% of text would convert, it's likely a mistake
        changed_chars = sum(1 for a, b in zip(text, converted_text) if a != b)
        if len(text) > 0:
            percentage = (changed_chars / len(text)) * 100
            if percentage > 50:
                logger.warning(
                    f"Preventing conversion with extremely high impact ({percentage:.2f}%)"
                )
                return text, False

        # Log that we did a conversion
        logger.info("Converted Symbol font Greek characters to Unicode")

        return converted_text, True

    def import_document(self, file_path: Union[str, Path]) -> Optional[Document]:
        """Import a document from a file.

        This copies the file to the storage directory and extracts metadata.

        Args:
            file_path: Path to the document file

        Returns:
            Imported document if successful, None otherwise
        """
        file_path = Path(file_path)

        if not file_path.exists():
            logger.error(f"File not found: {file_path}")
            return None

        try:
            # Create document object with metadata
            document = Document.from_file(file_path)

            # Copy file to storage directory
            storage_path = self.storage_dir / f"{document.id}{file_path.suffix}"
            shutil.copy2(file_path, storage_path)

            # Update file path to storage location
            document.file_path = Path(str(storage_path))

            # Extract text content if possible
            self.extract_text(document)

            # Save document to repository
            self.document_repository.save(document)

            return document
        except Exception as e:
            logger.error(f"Error importing document: {e}")
            return None

    def batch_import_documents(
        self,
        file_paths: List[Union[str, Path]],
        max_workers: int = 4,
        category_id: Optional[str] = None,
    ) -> List[Document]:
        """Import multiple documents at once.

        Args:
            file_paths: List of paths to document files
            max_workers: Maximum number of worker threads for parallel processing
            category_id: Optional category ID to assign to all imported documents

        Returns:
            List of successfully imported documents
        """
        logger.info(f"Batch importing {len(file_paths)} documents")
        imported_documents = []

        # Use ThreadPoolExecutor for parallel processing
        with concurrent.futures.ThreadPoolExecutor(max_workers=max_workers) as executor:
            # Submit all import tasks
            future_to_path = {
                executor.submit(self.import_document, path): path for path in file_paths
            }

            # Process results as they complete
            for future in concurrent.futures.as_completed(future_to_path):
                path = future_to_path[future]
                try:
                    document = future.result()
                    if document:
                        # Assign category if provided
                        if category_id is not None:
                            document.category = category_id
                            self.document_repository.save(document)

                        imported_documents.append(document)
                        logger.info(f"Successfully imported: {path}")
                    else:
                        logger.warning(f"Failed to import: {path}")
                except Exception as e:
                    logger.error(f"Exception importing {path}: {e}")

        logger.info(
            f"Batch import completed. {len(imported_documents)} of {len(file_paths)} documents imported successfully"
        )
        return imported_documents

    def import_documents_from_directory(
        self,
        directory_path: Union[str, Path],
        file_patterns: Optional[List[str]] = None,
        recursive: bool = False,
        max_workers: int = 4,
        category_id: Optional[str] = None,
    ) -> List[Document]:
        """Import all documents from a directory.

        Args:
            directory_path: Path to directory containing documents
            file_patterns: List of glob patterns to match files (e.g., ["*.pdf", "*.docx"])
            recursive: Whether to search subdirectories recursively
            max_workers: Maximum number of worker threads for parallel processing
            category_id: Optional category ID to assign to all imported documents

        Returns:
            List of successfully imported documents
        """
        directory_path = Path(directory_path)

        if not directory_path.exists() or not directory_path.is_dir():
            logger.error(f"Directory not found: {directory_path}")
            return []

        # Default patterns include all supported document types
        if file_patterns is None:
            file_patterns = [f"*.{t.value}" for t in DocumentType]

        # Find all matching files
        all_files = []
        for pattern in file_patterns:
            if recursive:
                # Search recursively through subdirectories
                search_pattern = str(directory_path / "**" / pattern)
                matches = glob(search_pattern, recursive=True)
            else:
                # Search only in the specified directory
                search_pattern = str(directory_path / pattern)
                matches = glob(search_pattern)

            all_files.extend(matches)

        # Remove duplicates
        all_files = list(set(all_files))

        logger.info(f"Found {len(all_files)} files to import in {directory_path}")

        # Convert string paths to Path objects before calling batch_import_documents
        # This fixes the type compatibility issue
        all_files_paths: List[Union[str, Path]] = [
            Path(f) if isinstance(f, str) else f for f in all_files
        ]
        return self.batch_import_documents(
            all_files_paths, max_workers=max_workers, category_id=category_id
        )

    def extract_text(self, document: Document) -> Optional[Document]:
        """Extract text content from a document.

        Args:
            document: Document to process

        Returns:
            Updated document if successful, None otherwise
        """
        if not os.path.exists(document.file_path):
            logger.error(f"Document file not found: {document.file_path}")
            return None

        try:
            # Extract text based on file type
            success = False
            if document.file_type == DocumentType.PDF:
                success = self._extract_text_from_pdf(document)
            elif document.file_type == DocumentType.DOCX:
                success = self._extract_text_from_docx(document)
            elif document.file_type == DocumentType.TXT:
                success = self._extract_text_from_txt(document)
            elif document.file_type.value == self.ODT:
                success = self._extract_text_from_odt(document)
            elif document.file_type.value == self.ODS:
                success = self._extract_text_from_ods(document)
            elif document.file_type.value == self.ODP:
                success = self._extract_text_from_odp(document)
            else:
                # If we get here, the file type is not supported
                logger.warning(
                    f"Text extraction not supported for {document.file_type}"
                )
                return None

            return document if success else None

        except Exception as e:
            logger.error(f"Error extracting text: {e}")
            return None

    def _remove_page_numbers(self, text: str) -> str:
        """Remove page numbers from extracted text.

        This function identifies and removes common page number patterns:
        - Standalone numbers at the beginning or end of a line
        - Numbers with "Page" prefix (e.g., "Page 1", "Page 2")
        - Roman numerals (i, ii, iii, iv, etc.) at line boundaries

        Args:
            text: The text to process

        Returns:
            Text with page numbers removed
        """
        # Pattern 1: Standalone numbers at beginning or end of a line
        text = re.sub(r"^\s*\d+\s*$", "", text, flags=re.MULTILINE)
        text = re.sub(r"^\s*\d+\s*\n", "\n", text, flags=re.MULTILINE)
        text = re.sub(r"\n\s*\d+\s*$", "\n", text, flags=re.MULTILINE)

        # Pattern 2: "Page X" or "Page X of Y" patterns
        text = re.sub(
            r"^\s*Page\s+\d+(\s+of\s+\d+)?\s*$",
            "",
            text,
            flags=re.MULTILINE | re.IGNORECASE,
        )
        text = re.sub(
            r"\n\s*Page\s+\d+(\s+of\s+\d+)?\s*\n",
            "\n",
            text,
            flags=re.MULTILINE | re.IGNORECASE,
        )

        # Pattern 3: Roman numerals at line boundaries
        # Match i, ii, iii, iv, v, vi, vii, viii, ix, x, etc.
        roman_pattern = r"^\s*[ivxlcdm]+\s*$"
        text = re.sub(roman_pattern, "", text, flags=re.MULTILINE | re.IGNORECASE)

        # Replace single line breaks with double line breaks where page numbers were removed
        # This preserves the spacing in the original document
        text = re.sub(r"\n\n\n+", "\n\n", text)

        return text

    def _extract_text_from_pdf(
        self, document: Document, remove_page_numbers: bool = True
    ) -> bool:
        """Extract text from a PDF document.

        Args:
            document: PDF document
            remove_page_numbers: Whether to remove page numbers from the extracted text (default: True)

        Returns:
            True if successful, False otherwise
        """
        try:
            pdf = fitz.open(document.file_path)
            text = ""
            page_count = 0
            word_count = 0

            # Extract text from each page
            for page in pdf:
                # Use the "text" option which gives plain text without page numbers in headers/footers
                # The default PyMuPDF text extraction preserves whitespace and ligatures
                page_text = page.get_text()

                # Convert any Symbol font Greek characters to Unicode
                page_text, was_converted = self._convert_symbol_to_greek(page_text)

                text += page_text
                page_count += 1
                # Approximate word count
                word_count += len(page_text.split())

            # Remove page numbers if requested
            if remove_page_numbers:
                text = self._remove_page_numbers(text)
                document.metadata["page_numbers_removed"] = True

            # Update document with extracted text and metadata
            document.content = text
            document.page_count = page_count
            document.word_count = word_count
            document.metadata["page_count"] = page_count
            document.metadata["word_count"] = word_count

            # Save updated document
            self.document_repository.save(document)

            return True
        except Exception as e:
            logger.error(f"Error extracting PDF text: {e}")
            return False

    def _extract_text_from_docx(self, document: Document) -> bool:
        """Extract text from a DOCX document with improved formatting preservation.

        Args:
            document: Document to extract text from

        Returns:
            True if successful, False otherwise
        """
        try:
            doc = DocxDocument(str(document.file_path))
            text_parts = []
            word_count = 0

            # Initialize metadata if needed
            document.metadata = document.metadata or {}
            document.metadata["extraction_status"] = "IN_PROGRESS"

            # Check for Symbol font usage in the document
            has_symbol_font = False

            # Process paragraphs with better formatting preservation
            for para in doc.paragraphs:
                para_text = ""

                # Process each run separately to handle font-specific conversions
                for run in para.runs:
                    run_text = run.text

                    # Only convert text that actually uses Symbol font
                    try:
                        if run.font.name == "Symbol" or "Symbol" in (
                            run.font.name or ""
                        ):
                            has_symbol_font = True
                            run_text, was_converted = self._convert_symbol_to_greek(
                                run_text
                            )
                            logger.debug(
                                f"Converted Symbol font text: '{run.text}' -> '{run_text}'"
                            )
                    except AttributeError:
                        # Some runs might not have font information
                        pass

                    para_text += run_text

                # Skip empty paragraphs but preserve line breaks
                if not para_text.strip():
                    text_parts.append("")  # Empty line to preserve paragraph breaks
                    continue

                # Add the paragraph with proper indentation for headers
                try:
                    if (
                        para.style
                        and para.style.name
                        and para.style.name.startswith("Heading")
                    ):
                        # Extract heading level from style name (e.g., 'Heading 1' -> level 1)
                        try:
                            level = int(para.style.name.split()[-1])
                            # Add appropriate marking for headings
                            prefix = "#" * level + " " if level <= 6 else ""
                            text_parts.append(f"{prefix}{para_text}")
                        except (ValueError, IndexError):
                            # Fallback if we can't extract level
                            text_parts.append(para_text)
                    else:
                        text_parts.append(para_text)
                except AttributeError:
                    # Fallback if style information is unavailable
                    text_parts.append(para_text)

                # Update word count
                word_count += len(para_text.split())

            # Process tables
            for table in doc.tables:
                text_parts.append("\n[TABLE]")
                for i, row in enumerate(table.rows):
                    row_texts = []
                    for cell in row.cells:
                        # Process cell paragraphs to handle Symbol font correctly
                        cell_text = ""
                        for cell_para in cell.paragraphs:
                            for cell_run in cell_para.runs:
                                run_text = cell_run.text

                                # Only convert text that actually uses Symbol font
                                try:
                                    if cell_run.font.name == "Symbol" or "Symbol" in (
                                        cell_run.font.name or ""
                                    ):
                                        has_symbol_font = True
                                        (
                                            run_text,
                                            was_converted,
                                        ) = self._convert_symbol_to_greek(run_text)
                                except AttributeError:
                                    pass

                                cell_text += run_text

                            # Add paragraph break within cell
                            cell_text += "\n"

                        # Clean cell text (remove excessive whitespace)
                        cell_text = " ".join(cell_text.split())
                        row_texts.append(cell_text)
                        word_count += len(cell_text.split())

                    # Add row with cell separator
                    text_parts.append(" | ".join(row_texts))

                    # Add separator line after header row
                    if i == 0:
                        text_parts.append(
                            "-" * sum(len(text) + 3 for text in row_texts)
                        )
                text_parts.append("[/TABLE]\n")

            # Join all parts with proper paragraph breaks
            full_text = "\n\n".join(text_parts)

            # Update document with extracted text
            document.content = full_text
            document.word_count = word_count
            document.metadata["extraction_status"] = "COMPLETED"
            document.metadata["has_greek_text"] = has_symbol_font
            document.metadata[
                "formatted_text"
            ] = True  # Indicate that this text preserves formatting

            # Save the document
            self.document_repository.save(document)

            return True

        except Exception as e:
            logger.error(f"Error extracting text from DOCX: {e}")
            document.metadata["extraction_status"] = "FAILED"
            document.metadata["extraction_error"] = str(e)
            return False

    def _detect_file_encoding(self, file_path: Path) -> str:
        """Detect the encoding of a text file using multiple methods.

        Args:
            file_path: Path to the text file

        Returns:
            Detected encoding name
        """
        import chardet

        # Read a sample of the file for detection
        with open(file_path, "rb") as file:
            raw_data = file.read(10000)  # Read first 10KB for detection

        # Use chardet for automatic detection
        detection_result = chardet.detect(raw_data)
        detected_encoding = detection_result.get("encoding", "utf-8")
        confidence = detection_result.get("confidence", 0.0)

        logger.info(
            f"Detected encoding: {detected_encoding} (confidence: {confidence:.2f})"
        )

        # If confidence is low, try common encodings for Word/LibreOffice files
        if confidence < 0.7:
            common_encodings = [
                "windows-1252",  # Common for Windows Word files
                "iso-8859-1",  # Latin-1, common for older files
                "cp1252",  # Windows Western European
                "utf-8-sig",  # UTF-8 with BOM
                "utf-16",  # UTF-16 with BOM
                "utf-16le",  # UTF-16 Little Endian
                "utf-16be",  # UTF-16 Big Endian
            ]

            for encoding in common_encodings:
                try:
                    with open(file_path, "r", encoding=encoding) as test_file:
                        test_file.read(1000)  # Try to read a portion
                    logger.info(f"Successfully validated encoding: {encoding}")
                    return encoding
                except (UnicodeDecodeError, UnicodeError):
                    continue

        return detected_encoding or "utf-8"

    def _convert_file_to_utf8(self, file_path: Path, source_encoding: str) -> bool:
        """Convert a text file from source encoding to UTF-8.

        Args:
            file_path: Path to the text file
            source_encoding: Source encoding of the file

        Returns:
            True if conversion was successful, False otherwise
        """
        try:
            # Create backup file path
            backup_path = file_path.with_suffix(file_path.suffix + ".backup")

            # Read file with source encoding
            with open(
                file_path, "r", encoding=source_encoding, errors="replace"
            ) as source_file:
                content = source_file.read()

            # Create backup of original file
            import shutil

            shutil.copy2(file_path, backup_path)

            # Write file with UTF-8 encoding
            with open(file_path, "w", encoding="utf-8") as target_file:
                target_file.write(content)

            logger.info(
                f"Successfully converted {file_path} from {source_encoding} to UTF-8"
            )
            logger.info(f"Backup created at {backup_path}")

            return True

        except Exception as e:
            logger.error(f"Error converting file {file_path} to UTF-8: {e}")
            return False

    def _extract_text_from_txt(self, document: Document) -> bool:
        """Extract text from a TXT document with robust encoding detection and conversion.

        Args:
            document: TXT document

        Returns:
            True if successful, False otherwise
        """
        try:
            file_path = Path(document.file_path)
            text = None
            encoding_used = None
            was_converted_to_utf8 = False

            # First, try to read as UTF-8
            try:
                with open(file_path, "r", encoding="utf-8") as file:
                    text = file.read()
                encoding_used = "utf-8"
                logger.info(f"Successfully read {file_path} as UTF-8")

            except UnicodeDecodeError as e:
                logger.warning(f"UTF-8 decode failed for {file_path}: {e}")

                # Detect the actual encoding
                detected_encoding = self._detect_file_encoding(file_path)
                logger.info(
                    f"Attempting to read with detected encoding: {detected_encoding}"
                )

                try:
                    # Try reading with detected encoding
                    with open(
                        file_path, "r", encoding=detected_encoding, errors="replace"
                    ) as file:
                        text = file.read()
                    encoding_used = detected_encoding

                    # Ask user if they want to convert the file to UTF-8
                    logger.info(
                        f"Successfully read {file_path} with encoding {detected_encoding}"
                    )

                    # Automatically convert to UTF-8 for future compatibility
                    if detected_encoding.lower() != "utf-8":
                        logger.info(
                            f"Converting {file_path} to UTF-8 for future compatibility"
                        )
                        was_converted_to_utf8 = self._convert_file_to_utf8(
                            file_path, detected_encoding
                        )

                except UnicodeDecodeError:
                    # If detected encoding fails, try common encodings with error handling
                    fallback_encodings = [
                        "windows-1252",
                        "iso-8859-1",
                        "cp1252",
                        "latin1",
                        "ascii",
                    ]

                    for fallback_encoding in fallback_encodings:
                        try:
                            with open(
                                file_path,
                                "r",
                                encoding=fallback_encoding,
                                errors="replace",
                            ) as file:
                                text = file.read()
                            encoding_used = fallback_encoding
                            logger.warning(
                                f"Used fallback encoding {fallback_encoding} for {file_path}"
                            )

                            # Convert to UTF-8
                            was_converted_to_utf8 = self._convert_file_to_utf8(
                                file_path, fallback_encoding
                            )
                            break

                        except UnicodeDecodeError:
                            continue

                    if text is None:
                        # Last resort: read as binary and decode with errors='replace'
                        with open(file_path, "rb") as file:
                            raw_data = file.read()
                        text = raw_data.decode("utf-8", errors="replace")
                        encoding_used = "utf-8 (with errors replaced)"
                        logger.warning(f"Used error replacement for {file_path}")

            if text is None:
                raise Exception("Could not read file with any encoding")

            # Convert potential Symbol font Greek text
            text, was_symbol_converted = self._convert_symbol_to_greek(text)

            # Update document with text content
            document.content = text
            document.word_count = len(text.split())
            document.page_count = 1  # Plain text files don't have pages
            document.metadata["word_count"] = len(text.split())
            document.metadata["page_count"] = 1
            document.metadata["original_encoding"] = encoding_used
            document.metadata["converted_to_utf8"] = was_converted_to_utf8
            document.metadata["symbol_conversion"] = was_symbol_converted

            # Save updated document
            self.document_repository.save(document)

            return True

        except Exception as e:
            logger.error(f"Error extracting TXT text: {e}")
            return False

    def _extract_text_from_odt(self, document: Document) -> bool:
        """Extract text from a LibreOffice Writer (ODT) document.

        Args:
            document: ODT document

        Returns:
            True if successful, False otherwise
        """
        try:
            # Attempt to use odfpy library for extraction
            from odf import style, teletype, text
            from odf.opendocument import load

            textdoc = load(document.file_path)
            allparas = textdoc.getElementsByType(text.P)

            extracted_text = ""
            word_count = 0

            # Check for Symbol font usage
            has_symbol_font = False
            styles = textdoc.getElementsByType(style.Style)
            for s in styles:
                text_props = s.getElementsByType(style.TextProperties)
                for prop in text_props:
                    if prop.getAttribute("fontname") and "Symbol" in prop.getAttribute(
                        "fontname"
                    ):
                        has_symbol_font = True
                        break
                if has_symbol_font:
                    break

            for para in allparas:
                para_text = teletype.extractText(para)

                # Check for Greek letters if Symbol font is used
                if has_symbol_font:
                    para_text, was_converted = self._convert_symbol_to_greek(para_text)

                extracted_text += para_text + "\n"
                word_count += len(para_text.split())

            # Approximate page count (1 page per 500 words)
            page_count = max(1, word_count // 500)

            # Update document with text content
            document.content = extracted_text
            document.word_count = word_count
            document.page_count = page_count
            document.metadata["word_count"] = word_count
            document.metadata["page_count"] = page_count
            document.metadata["has_greek_text"] = has_symbol_font

            # Save updated document
            self.document_repository.save(document)

            return was_converted
        except ImportError:
            logger.error(
                "odfpy library not installed. Please install odfpy to process ODT files."
            )
            return False
        except Exception as e:
            logger.error(f"Error extracting ODT text: {e}")
            return False

    def _extract_text_from_ods(self, document: Document) -> bool:
        """Extract text from a LibreOffice Calc (ODS) document.

        Args:
            document: ODS document

        Returns:
            True if successful, False otherwise
        """
        try:
            # Attempt to use odfpy library for extraction
            from odf import table, teletype, text
            from odf.opendocument import load

            spreadsheet = load(document.file_path)

            # Get all tables (sheets)
            tables = spreadsheet.getElementsByType(table.Table)

            extracted_text = ""
            word_count = 0
            sheet_count = len(tables)

            # Process each table (sheet)
            for sheet in tables:
                sheet_name = sheet.getAttribute("name")
                extracted_text += f"Sheet: {sheet_name}\n\n"

                # Get all rows
                rows = sheet.getElementsByType(table.TableRow)
                for row in rows:
                    row_text = ""
                    # Get all cells in this row
                    cells = row.getElementsByType(table.TableCell)
                    for cell in cells:
                        # Get cell content (which may include paragraphs)
                        paras = cell.getElementsByType(text.P)
                        for para in paras:
                            cell_text = teletype.extractText(para)
                            row_text += cell_text + "\t"

                    extracted_text += row_text + "\n"
                    word_count += len(row_text.split())

                extracted_text += "\n\n"

            # Update document with text content
            document.content = extracted_text
            document.word_count = word_count
            document.page_count = sheet_count
            document.metadata["word_count"] = word_count
            document.metadata["sheet_count"] = sheet_count
            document.metadata["page_count"] = sheet_count

            # Save updated document
            self.document_repository.save(document)

            return True
        except ImportError:
            logger.error(
                "odfpy library not installed. Please install odfpy to process ODS files."
            )
            return False
        except Exception as e:
            logger.error(f"Error extracting ODS text: {e}")
            return False

    def _extract_text_from_odp(self, document: Document) -> bool:
        """Extract text from a LibreOffice Impress (ODP) document.

        Args:
            document: ODP document

        Returns:
            True if successful, False otherwise
        """
        try:
            # Attempt to use odfpy library for extraction
            from odf import draw, teletype, text
            from odf.opendocument import load

            presentation = load(document.file_path)

            # Get all slides (pages)
            slides = presentation.getElementsByType(draw.Page)

            extracted_text = ""
            word_count = 0
            slide_count = len(slides)

            # Process each slide
            for slide_idx, slide in enumerate(slides, 1):
                slide_name = slide.getAttribute("name") or f"Slide {slide_idx}"
                extracted_text += f"Slide: {slide_name}\n\n"

                # Get all text frames in this slide
                frames = slide.getElementsByType(draw.Frame)
                for frame in frames:
                    # Get text boxes in this frame
                    textboxes = frame.getElementsByType(draw.TextBox)
                    for textbox in textboxes:
                        # Get paragraphs in this text box
                        paras = textbox.getElementsByType(text.P)
                        for para in paras:
                            para_text = teletype.extractText(para)
                            extracted_text += para_text + "\n"
                            word_count += len(para_text.split())

                extracted_text += "\n\n"

            # Update document with text content
            document.content = extracted_text
            document.word_count = word_count
            document.page_count = slide_count
            document.metadata["word_count"] = word_count
            document.metadata["slide_count"] = slide_count
            document.metadata["page_count"] = slide_count

            # Save updated document
            self.document_repository.save(document)

            return True
        except ImportError:
            logger.error(
                "odfpy library not installed. Please install odfpy to process ODP files."
            )
            return False
        except Exception as e:
            logger.error(f"Error extracting ODP text: {e}")
            return False

    def get_document(self, document_id: str) -> Optional[Document]:
        """Get a document by ID with integrity checks.

        Args:
            document_id: Document ID

        Returns:
            Document if found and valid, None otherwise
        """
        document = self.document_repository.get_by_id(document_id)

        if not document:
            logger.warning(f"Document not found: {document_id}")
            return None

        # Perform integrity checks
        document.metadata = document.metadata or {}
        document.metadata["integrity_status"] = "OK"

        # Verify file integrity for file-based documents
        if document.file_path:
            # Check if file exists
            file_path = Path(document.file_path)
            if not file_path.exists():
                logger.warning(f"Document file missing: {document.file_path}")
                document.metadata["integrity_status"] = "FILE_MISSING"
                return document

            # Check if file size matches recorded size
            if document.size_bytes:
                # For QGemDocuments, size can change due to HTML formatting - don't warn
                if document.file_type == DocumentType.QTDOC:
                    # Update size silently without warnings
                    if hasattr(document, "html_content") and document.html_content:
                        document.size_bytes = len(document.html_content.encode("utf-8"))
                elif document.file_path.exists():
                    actual_size = file_path.stat().st_size
                    if (
                        abs(actual_size - document.size_bytes) > 10
                    ):  # Allow small differences
                        logger.warning(
                            f"Document size mismatch for {document_id}: "
                            f"recorded={document.size_bytes}, actual={actual_size}"
                        )
                        document.metadata["integrity_status"] = "SIZE_MISMATCH"

            # Additional check: verify checksum if available
            if document.metadata.get("checksum"):
                actual_checksum = self._calculate_file_checksum(file_path)
                if actual_checksum != document.metadata["checksum"]:
                    logger.warning(
                        f"Document checksum mismatch for {document_id}: "
                        f"recorded={document.metadata['checksum']}, actual={actual_checksum}"
                    )
                    document.metadata["integrity_status"] = "CHECKSUM_MISMATCH"
                    return document

        return document

    def _calculate_file_checksum(self, file_path: Path) -> str:
        """Calculate MD5 checksum of a file.

        Args:
            file_path: Path to the file

        Returns:
            MD5 checksum as a hexadecimal string
        """
        import hashlib

        md5 = hashlib.md5()
        with open(file_path, "rb") as f:
            # Read in chunks to avoid loading large files entirely into memory
            for chunk in iter(lambda: f.read(4096), b""):
                md5.update(chunk)
        return md5.hexdigest()

    def get_all_documents(self) -> List[Document]:
        """Get all documents from the repository.

        Returns:
            List of all documents
        """
        return self.document_repository.get_all()

    def list_documents(self) -> List[Document]:
        """List all documents from the repository.

        This is an alias for get_all_documents() to maintain compatibility
        with different parts of the application.

        Returns:
            List of all documents
        """
        return self.get_all_documents()

    def _validate_document(self, document: Document) -> Tuple[bool, Optional[str]]:
        """Validate document before saving to database.

        Performs checks on the document to ensure it meets requirements:
        - File exists (for non-memory documents)
        - Document size is within reasonable limits
        - Required metadata is present

        Args:
            document: Document to validate

        Returns:
            Tuple of (is_valid, error_message)
        """
        # Check if file exists (for non-memory documents)
        if document.file_path and not os.path.exists(document.file_path):
            return False, f"File does not exist: {document.file_path}"

        # Check if document has a name
        if not document.name or document.name.strip() == "":
            return False, "Document must have a name"

        # Check if extraction status indicates a failure
        if document.metadata.get("extraction_status") == "FAILED":
            return False, "Text extraction failed for this document"

        # Check document size (if applicable)
        if document.size_bytes and document.size_bytes > 100 * 1024 * 1024:  # 100 MB
            return (
                False,
                f"Document size ({document.size_bytes / 1024 / 1024:.2f} MB) exceeds maximum allowed (100 MB)",
            )

        # Check content length if extracted
        if (
            document.content and len(document.content) > 10 * 1024 * 1024
        ):  # 10 MB of text
            return False, "Extracted text size exceeds maximum allowed (10 MB)"

        # All validation passed
        return True, None

    def save_document(self, document: Document) -> Optional[Document]:
        """Save a document.

        Args:
            document: Document to save

        Returns:
            Saved document if successful, None otherwise
        """
        # Validate document before saving
        is_valid, error_message = self._validate_document(document)
        if not is_valid:
            logger.error(f"Document validation failed: {error_message}")
            return None

        # Calculate and store checksum for future integrity checks
        if document.file_path:
            document.metadata = document.metadata or {}
            document.metadata["checksum"] = self._calculate_file_checksum(
                Path(document.file_path)
            )

        # Proceed with saving
        success = self.document_repository.save(document)
        if success:
            return document
        return None

    def delete_document(self, document_id: str) -> bool:
        """Delete a document.

        Args:
            document_id: ID of the document to delete

        Returns:
            True if deleted successfully, False otherwise
        """
        try:
            # Get the document
            document = self.get_document(document_id)
            if not document:
                logger.error(f"Document with ID {document_id} not found")
                return False

            # Delete from repository
            success = self.document_repository.delete(document_id)
            if not success:
                logger.error(
                    f"Failed to delete document with ID {document_id} from repository"
                )
                return False

            # Delete associated file if it exists
            if document.file_path and os.path.exists(document.file_path):
                try:
                    os.remove(document.file_path)
                except Exception as e:
                    logger.warning(f"Error deleting file: {e}")
                    # Still return True as document was deleted from repository

            logger.info(f"Document {document_id} deleted successfully")
            return True
        except Exception as e:
            logger.error(f"Error deleting document: {e}")
            return False

    def search_documents(
        self,
        query: str = "",
        category: Optional[str] = None,
        doc_type: Optional[DocumentType] = None,
        tag_ids: Optional[List[str]] = None,
        date_from: Optional[datetime] = None,
        date_to: Optional[datetime] = None,
    ) -> List[Document]:
        """Search for documents with various filters.

        Args:
            query: Text search query
            category: Filter by category
            doc_type: Filter by document type
            tag_ids: Filter by tag IDs
            date_from: Filter by creation date (from)
            date_to: Filter by creation date (to)

        Returns:
            List of matching documents
        """
        # Get documents and total count
        documents, _ = self.document_repository.search(
            query=query,
            category=category,
            doc_type=doc_type,
            tags=tag_ids,
            date_from=date_from,
            date_to=date_to,
        )

        return documents

    def add_tag_to_document(self, document_id: str, tag_id: str) -> bool:
        """Add a tag to a document.

        Args:
            document_id: Document ID
            tag_id: Tag ID

        Returns:
            True if successful, False otherwise
        """
        document = self.document_repository.get_by_id(document_id)

        if not document:
            logger.error(f"Document not found: {document_id}")
            return False

        document.add_tag(tag_id)
        return self.document_repository.save(document)

    def remove_tag_from_document(self, document_id: str, tag_id: str) -> bool:
        """Remove a tag from a document.

        Args:
            document_id: Document ID
            tag_id: Tag ID

        Returns:
            True if successful, False otherwise
        """
        document = self.document_repository.get_by_id(document_id)

        if not document:
            logger.error(f"Document not found: {document_id}")
            return False

        document.remove_tag(tag_id)
        return self.document_repository.save(document)

    def set_document_category(
        self, document_id: str, category_id: Optional[str]
    ) -> bool:
        """Set the category for a document.

        Args:
            document_id: Document ID
            category_id: Category ID, or None to clear

        Returns:
            True if successful, False otherwise
        """
        document = self.document_repository.get_by_id(document_id)

        if not document:
            logger.error(f"Document not found: {document_id}")
            return False

        document.category = category_id
        return self.document_repository.save(document)

    def get_documents_by_category(self, category_id: str) -> List[Document]:
        """Get documents by category.

        Args:
            category_id: Category ID

        Returns:
            List of documents in the category
        """
        return self.document_repository.get_by_category(category_id)

    def update_parsed_text(self, document_id: str, new_text: str) -> bool:
        """Update the extracted plain text content of a document.

        Args:
            document_id: The ID of the document to update.
            new_text: The new plain text content.

        Returns:
            True if the update was successful, False otherwise.
        """
        logger.debug(f"Attempting to update parsed text for document ID: {document_id}")
        document = self.document_repository.get_by_id(document_id)

        if not document:
            logger.warning(
                f"Document not found with ID: {document_id} for text update."
            )
            return False

        document.extracted_text = new_text
        document.last_modified_date = datetime.now()
        # Potentially update word count if that's dynamically calculated or stored
        # For now, just updating text and modification date.

        try:
            if self.document_repository.save(document):
                logger.info(
                    f"Successfully updated parsed text for document ID: {document_id}"
                )
                return True
            else:
                logger.error(
                    f"Failed to save updated parsed text for document ID: {document_id} via repository."
                )
                return False
        except Exception as e:
            logger.error(
                f"Exception while saving updated parsed text for document ID {document_id}: {e}"
            )
            return False

    def revert_greek_conversion(self, document_id: str) -> bool:
        """Revert a document that was incorrectly converted to Greek back to original text.

        This is a recovery function for documents that were incorrectly identified as
        containing Symbol font Greek text.

        Args:
            document_id: ID of document to revert

        Returns:
            True if successful, False otherwise
        """
        document = self.get_document(document_id)
        if not document or not document.content:
            logger.error(
                f"Cannot revert document {document_id}: document not found or has no content"
            )
            return False

        try:
            # Check if document was flagged as having Greek text
            if not document.metadata.get("has_greek_text"):
                logger.warning(
                    f"Document {document_id} was not marked as having Greek text"
                )
                return False

            # Create a reverse mapping from Greek to Latin
            reverse_mapping = {
                greek: latin for latin, greek in self.SYMBOL_TO_GREEK.items()
            }
            reverse_mapping.update(
                {
                    replacement: symbol
                    for symbol, replacement in self.SYMBOL_SPECIAL.items()
                }
            )

            # Create a new text with the conversion reversed
            original_text = document.content
            reverted_text = original_text

            # Apply the reverse mapping
            for greek, latin in reverse_mapping.items():
                reverted_text = reverted_text.replace(greek, latin)

            # Update the document
            document.content = reverted_text
            document.metadata["has_greek_text"] = False
            document.metadata["greek_conversion_reverted"] = True
            document.metadata["conversion_reverted_date"] = datetime.now().isoformat()

            # Save the updated document
            self.document_repository.save(document)

            # Log the change
            logger.info(
                f"Successfully reverted Greek conversion for document {document_id}"
            )
            return True

        except Exception as e:
            logger.error(f"Error reverting Greek conversion: {e}")
            return False

    def _analyze_conversion_impact(
        self, original_text: str, converted_text: str
    ) -> None:
        """Analyze and log the impact of Greek text conversion.

        Args:
            original_text: Text before conversion
            converted_text: Text after conversion
        """
        if not original_text or not converted_text or original_text == converted_text:
            return

        # Count characters that were changed
        changed_chars = sum(1 for a, b in zip(original_text, converted_text) if a != b)

        # Calculate percentage of text affected
        total_chars = len(original_text)
        if total_chars > 0:
            percentage = (changed_chars / total_chars) * 100

            # Log the impact
            if percentage > 20:
                logger.warning(
                    f"High impact conversion: {percentage:.2f}% of text was converted to Greek"
                )
            else:
                logger.info(
                    f"Greek conversion affected {percentage:.2f}% of document text"
                )

            # If more than 50% was converted, it's likely a mistake
            if percentage > 50:
                logger.warning(
                    "Suspiciously high conversion rate - might be incorrect identification of Symbol font"
                )

    def update_document(self, document: Document) -> bool:
        """Update an existing document in the repository.

        Args:
            document: Document object to update

        Returns:
            True if update was successful, False otherwise
        """
        logger.debug(f"Updating document: {document.id}")

        # Validate document
        is_valid, error_message = self._validate_document(document)
        if not is_valid:
            logger.error(f"Document validation failed: {error_message}")
            return False

        # Update last_modified_date to current time
        document.last_modified_date = datetime.now()

        # Save to repository
        success = self.document_repository.save(document)

        if success:
            logger.info(f"Document {document.id} updated successfully")
            return True
        else:
            logger.error(f"Failed to update document {document.id}")
            return False

    def bulk_convert_text_files_to_utf8(
        self,
        directory_path: Union[str, Path],
        file_patterns: Optional[List[str]] = None,
        recursive: bool = True,
    ) -> Dict[str, Any]:
        """Convert multiple text files in a directory to UTF-8 encoding.

        Args:
            directory_path: Directory containing text files to convert
            file_patterns: List of file patterns to match (e.g., ['*.txt', '*.text'])
            recursive: Whether to search subdirectories

        Returns:
            Dictionary with conversion results and statistics
        """
        directory_path = Path(directory_path)

        if file_patterns is None:
            file_patterns = ["*.txt", "*.text", "*.asc"]

        results = {
            "total_files": 0,
            "converted_files": 0,
            "already_utf8": 0,
            "failed_conversions": 0,
            "conversion_details": [],
            "errors": [],
        }

        try:
            # Find all matching text files
            text_files = []
            for pattern in file_patterns:
                if recursive:
                    text_files.extend(directory_path.rglob(pattern))
                else:
                    text_files.extend(directory_path.glob(pattern))

            results["total_files"] = len(text_files)
            logger.info(f"Found {len(text_files)} text files to process")

            for file_path in text_files:
                try:
                    # Detect current encoding
                    detected_encoding = self._detect_file_encoding(file_path)

                    file_result = {
                        "file_path": str(file_path),
                        "original_encoding": detected_encoding,
                        "status": "unknown",
                    }

                    # Check if already UTF-8
                    if detected_encoding.lower() in ["utf-8", "ascii"]:
                        file_result["status"] = "already_utf8"
                        results["already_utf8"] += 1
                        logger.info(f"File {file_path} is already UTF-8 compatible")
                    else:
                        # Convert to UTF-8
                        if self._convert_file_to_utf8(file_path, detected_encoding):
                            file_result["status"] = "converted"
                            file_result["converted_to"] = "utf-8"
                            results["converted_files"] += 1
                            logger.info(
                                f"Successfully converted {file_path} from {detected_encoding} to UTF-8"
                            )
                        else:
                            file_result["status"] = "failed"
                            results["failed_conversions"] += 1
                            logger.error(f"Failed to convert {file_path}")

                    results["conversion_details"].append(file_result)

                except Exception as e:
                    error_msg = f"Error processing {file_path}: {e}"
                    results["errors"].append(error_msg)
                    results["failed_conversions"] += 1
                    logger.error(error_msg)

            # Log summary
            logger.info("Bulk conversion complete:")
            logger.info(f"  Total files: {results['total_files']}")
            logger.info(f"  Converted: {results['converted_files']}")
            logger.info(f"  Already UTF-8: {results['already_utf8']}")
            logger.info(f"  Failed: {results['failed_conversions']}")

            return results

        except Exception as e:
            error_msg = f"Error during bulk conversion: {e}"
            results["errors"].append(error_msg)
            logger.error(error_msg)
            return results
